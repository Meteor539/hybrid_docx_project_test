from typing import Any

import re
from docx import Document


class DocxStructureAdapter:
    """对现有 python-docx 解析链路的适配器。"""

    _statement_headings = [
        "原创性声明",
        "学位论文原创性声明",
        "学位论文版权使用授权书",
        "版权使用授权书",
        "使用授权书",
        "使用授权声明",
    ]

    def __init__(self) -> None:
        self.parser = None
        self.parts_order: list[str] = []

    def parse(self, file_path: str) -> tuple[Any, dict[str, Any]]:
        try:
            if self.parser is None:
                from model.document_parser import DocumentParser

                self.parser = DocumentParser()
            sections = self.parser.parse_document(file_path)
            self.parts_order = list(getattr(self.parser, "parts_order", []) or [])
            return self.parser.doc, sections
        except Exception:
            # 旧解析器不可用时，回退到纯 python-docx 解析。
            return self._fallback_parse(file_path)

    def _fallback_parse(self, file_path: str) -> tuple[Any, dict[str, Any]]:
        doc = Document(file_path)
        paragraphs = [p for p in doc.paragraphs if p.text and p.text.strip()]

        sections = self._empty_sections()
        self.parts_order = []

        if not paragraphs:
            return doc, sections

        # 粗粒度边界识别：声明、摘要、目录、正文、参考文献、致谢
        idx_statement = self._find_statement_index(paragraphs)
        idx_cn_abs = self._find_exact_heading_index(paragraphs, ["摘要", "摘 要"])
        idx_en_abs = self._find_exact_heading_index(paragraphs, ["abstract"], lower=True)
        idx_catalogue = self._find_exact_heading_index(paragraphs, ["目录"])
        idx_main = self._find_chapter_index(paragraphs)
        idx_ref = self._find_exact_heading_index(paragraphs, ["参考文献", "references", "bibliography"], lower=True)
        idx_ack = self._find_ack_index(paragraphs)

        # 封面
        cover_end = self._first_valid_index(
            [idx_statement, idx_cn_abs, idx_catalogue, idx_main, idx_ref, idx_ack],
            default=len(paragraphs),
        )
        cover = paragraphs[:cover_end]
        if cover:
            self.parts_order.append("cover")
            sections["cover"]["school"] = cover[0]
            if len(cover) > 1:
                sections["cover"]["title"] = cover[1]
            if len(cover) > 2:
                sections["cover"]["personal_information"] = cover[2:]

        # 声明（简单聚合）
        if idx_statement is not None:
            self.parts_order.append("statement1")
            st_end = self._first_index_after(
                idx_statement,
                [idx_cn_abs, idx_en_abs, idx_catalogue, idx_main, idx_ref, idx_ack],
                default=len(paragraphs),
            )
            st = paragraphs[idx_statement:st_end]
            if st:
                statement_titles, statement_contents = self._split_statement_block(st)
                sections["statement"]["title"] = statement_titles
                sections["statement"]["content"] = statement_contents

        # 中文摘要
        if idx_cn_abs is not None:
            self.parts_order.append("chinese_abstract")
            cn_end = self._first_index_after(
                idx_cn_abs,
                [idx_en_abs, idx_catalogue, idx_main, idx_ref, idx_ack],
                default=len(paragraphs),
            )
            cn = paragraphs[idx_cn_abs:cn_end]
            if cn:
                sections["abstract_keyword"]["chinese_title"] = cn[0]
                self._assign_abstract_parts(sections, cn, language="chinese")

        # 英文摘要
        if idx_en_abs is not None:
            self.parts_order.append("english_abstract")
            en_end = self._first_index_after(
                idx_en_abs,
                [idx_catalogue, idx_main, idx_ref, idx_ack],
                default=len(paragraphs),
            )
            en = paragraphs[idx_en_abs:en_end]
            if en:
                sections["abstract_keyword"]["english_title"] = en[0]
                self._assign_abstract_parts(sections, en, language="english")

        if idx_catalogue is not None:
            self.parts_order.append("catalogue")
            catalogue_end = self._first_index_after(
                idx_catalogue,
                [idx_main, idx_ref, idx_ack],
                default=len(paragraphs),
            )
            cat = paragraphs[idx_catalogue:catalogue_end]
            if cat:
                sections["catalogue"]["title"] = cat[0]
                if len(cat) > 1:
                    sections["catalogue"]["content"] = cat[1:]

        # 正文及标题分类
        main_start = idx_main if idx_main is not None else 0
        main_end = self._first_index_after(
            main_start,
            [idx_ref, idx_ack],
            default=len(paragraphs),
        )
        if main_start < main_end:
            self.parts_order.append("main_text")
            for para in paragraphs[main_start:main_end]:
                text = para.text.strip()
                if self._is_chapter_title(text):
                    sections["headings"]["chapter"].append(para)
                elif self._is_level2_title(text):
                    sections["headings"]["level2"].append(para)
                elif self._is_level1_title(text):
                    sections["headings"]["level1"].append(para)
                elif self._is_level3_title(text):
                    sections["headings"]["level3"].append(para)
                elif self._is_figure_or_table_title(text):
                    sections["figures_or_tables_title"].append(para)
                else:
                    sections["main_text"].append(para)

        # 参考文献
        if idx_ref is not None:
            self.parts_order.append("references")
            ref_end = self._first_index_after(idx_ref, [idx_ack], default=len(paragraphs))
            refs = paragraphs[idx_ref:ref_end]
            if refs:
                sections["references"]["title"] = refs[0]
                if len(refs) > 1:
                    sections["references"]["content"] = refs[1:]

        # 致谢
        if idx_ack is not None:
            self.parts_order.append("acknowledgments")
            ack = paragraphs[idx_ack:]
            if ack:
                sections["acknowledgments"]["title"] = ack[0]
                if len(ack) > 1:
                    sections["acknowledgments"]["content"] = ack[1:]

        sections["figures"] = doc.inline_shapes
        sections["tables"] = doc.tables
        return doc, sections

    @staticmethod
    def _assign_abstract_parts(sections: dict[str, Any], paragraphs, *, language: str) -> None:
        if not paragraphs:
            return

        content_key = f"{language}_content"
        keyword_title_key = f"{language}_keyword_title"
        keyword_key = f"{language}_keyword"

        body = paragraphs[1:]
        if not body:
            return

        keyword_idx = DocxStructureAdapter._find_keyword_paragraph_index(body, language=language)
        if keyword_idx is not None:
            keyword_paragraph = body[keyword_idx]
            content_paragraphs = body[:keyword_idx]
            sections["abstract_keyword"][content_key] = content_paragraphs or None
            sections["abstract_keyword"][keyword_title_key] = keyword_paragraph
            sections["abstract_keyword"][keyword_key] = keyword_paragraph
            return

        if len(body) >= 2:
            sections["abstract_keyword"][content_key] = body[:-1]
            sections["abstract_keyword"][keyword_key] = body[-1]
        else:
            sections["abstract_keyword"][keyword_key] = body[-1]

    @staticmethod
    def _find_keyword_paragraph_index(paragraphs, *, language: str):
        if language == "english":
            pattern = r"^\s*key\s*words?\s*[:：]"
        else:
            pattern = r"^\s*(关键词|关\s*键\s*词)\s*[:：]"

        for idx, para in enumerate(paragraphs):
            text = (para.text or "").strip()
            if re.match(pattern, text, flags=re.IGNORECASE):
                return idx
        return None

    @staticmethod
    def _empty_sections() -> dict[str, Any]:
        return {
            "cover": {"school": None, "title": None, "personal_information": None},
            "statement": {"title": [], "content": []},
            "abstract_keyword": {
                "chinese_title": None,
                "chinese_content": None,
                "chinese_keyword_title": None,
                "chinese_keyword": None,
                "english_title": None,
                "english_content": None,
                "english_keyword_title": None,
                "english_keyword": None,
            },
            "catalogue": {"title": None, "content": None},
            "main_text": [],
            "headings": {"chapter": [], "level1": [], "level2": [], "level3": []},
            "figures_or_tables_title": [],
            "figures": [],
            "tables": [],
            "references": {"title": None, "content": []},
            "acknowledgments": {"title": None, "content": None},
        }

    @staticmethod
    def _find_index(paragraphs, keywords: list[str], lower: bool = False):
        for idx, para in enumerate(paragraphs):
            text = para.text.strip()
            candidate = text.lower() if lower else text
            keys = [k.lower() for k in keywords] if lower else keywords
            if any(k in candidate for k in keys):
                return idx
        return None

    @staticmethod
    def _normalize_heading_text(text: str, *, lower: bool = False) -> str:
        normalized = re.sub(r"[\s\u3000]+", "", text or "")
        if lower:
            normalized = normalized.lower()
        return normalized

    @classmethod
    def _find_exact_heading_index(cls, paragraphs, headings: list[str], lower: bool = False):
        normalized_headings = {
            cls._normalize_heading_text(h, lower=lower)
            for h in headings
        }
        for idx, para in enumerate(paragraphs):
            normalized = cls._normalize_heading_text(para.text, lower=lower)
            if normalized in normalized_headings:
                return idx
        return None

    @classmethod
    def _find_statement_index(cls, paragraphs):
        return cls._find_exact_heading_index(paragraphs, cls._statement_headings)

    @classmethod
    def _is_statement_heading(cls, text: str) -> bool:
        normalized = cls._normalize_heading_text(text)
        normalized_headings = {cls._normalize_heading_text(item) for item in cls._statement_headings}
        return normalized in normalized_headings

    @classmethod
    def _split_statement_block(cls, paragraphs):
        titles = []
        contents = []
        for paragraph in paragraphs or []:
            text = (getattr(paragraph, "text", "") or "").strip()
            if not text:
                continue
            if cls._is_statement_heading(text):
                titles.append(paragraph)
            else:
                contents.append(paragraph)
        return titles, contents

    @classmethod
    def _find_ack_index(cls, paragraphs):
        for idx, para in enumerate(paragraphs):
            normalized = cls._normalize_heading_text(para.text, lower=True)
            if normalized in {"致谢", "acknowledgments", "acknowledgement"}:
                return idx
        return None

    @staticmethod
    def _find_chapter_index(paragraphs):
        for idx, para in enumerate(paragraphs):
            if DocxStructureAdapter._is_chapter_title(para.text.strip()):
                return idx
        return None

    @staticmethod
    def _first_valid_index(values: list[int | None], default: int) -> int:
        valid = [v for v in values if isinstance(v, int)]
        return min(valid) if valid else default

    @staticmethod
    def _first_index_after(base: int, values: list[int | None], default: int) -> int:
        valid = [v for v in values if isinstance(v, int) and v > base]
        return min(valid) if valid else default

    @staticmethod
    def _is_chapter_title(text: str) -> bool:
        if not text:
            return False
        normalized = text.strip()
        if re.search(r"[，。；：！？!?]", normalized):
            return False
        if len(normalized) > 40:
            return False
        return bool(
            re.fullmatch(r"第[一二三四五六七八九十百千0-9]+章\s*[^\s，。；：！？!?]{1,30}", normalized)
            or re.fullmatch(r"chapter\s+\d+\s+.{1,30}", normalized, flags=re.IGNORECASE)
        )

    @staticmethod
    def _is_level1_title(text: str) -> bool:
        if not text:
            return False
        normalized = text.strip()
        if re.search(r"[，。；：！？!?]", normalized):
            return False
        if len(normalized) > 40:
            return False
        return bool(
            re.fullmatch(r"\d+\.\d+\s*[^\s，。；：！？!?].{0,30}", normalized)
            or re.fullmatch(r"[一二三四五六七八九十]+、[^\s，。；：！？!?].{0,30}", normalized)
        )

    @staticmethod
    def _is_level2_title(text: str) -> bool:
        if not text:
            return False
        normalized = text.strip()
        if re.search(r"[，。；：！？!?]", normalized):
            return False
        if len(normalized) > 40:
            return False
        return bool(
            re.fullmatch(r"\d+\.\d+\.\d+\s*[^\s，。；：！？!?].{0,30}", normalized)
            or re.fullmatch(r"（[一二三四五六七八九十]+）[^\s，。；：！？!?].{0,30}", normalized)
        )

    @staticmethod
    def _is_level3_title(text: str) -> bool:
        if not text:
            return False
        normalized = text.strip()
        if re.search(r"[，。；：！？!?]", normalized):
            return False
        if len(normalized) > 30:
            return False
        return bool(
            re.fullmatch(r"\d+\.\s*[^\s，。；：！？!?].{0,20}", normalized)
            or re.fullmatch(r"\d+\)[^\s，。；：！？!?].{0,20}", normalized)
        )

    @staticmethod
    def _is_figure_or_table_title(text: str) -> bool:
        return bool(re.match(r"^(图|表)\s*\d+([\.．]\d+)*", text))
